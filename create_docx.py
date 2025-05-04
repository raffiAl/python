from docx import Document
from docx.shared import Pt

def create_document():
    # Membuat dokumen baru
    doc = Document()
    
    # Menambahkan judul
    doc.add_heading('Pembuatan Invers Matriks dengan Eliminasi Gauss-Jordan', 0)
    
    # Langkah 1: Menulis Matriks yang Diperluas dengan Matriks Identitas
    doc.add_heading('Langkah 1: Menulis Matriks yang Diperluas', level=1)
    doc.add_paragraph('Matriks asli \( A \):')
    table = doc.add_table(rows=3, cols=3)
    data = [
        [2, 0, 0],
        [1, 5, 0],
        [7, 9, 4]
    ]
    for i in range(3):
        row_cells = table.rows[i].cells
        for j in range(3):
            row_cells[j].text = str(data[i][j])
    
    doc.add_paragraph('Matriks yang diperluas dengan matriks identitas:')
    extended_table = doc.add_table(rows=3, cols=6)
    extended_data = [
        [2, 0, 0, 1, 0, 0],
        [1, 5, 0, 0, 1, 0],
        [7, 9, 4, 0, 0, 1]
    ]
    for i in range(3):
        row = extended_data[i]
        for j in range(6):
            extended_table.cell(i, j).text = str(row[j])
    
    # Langkah 2: Operasi Baris untuk Mengubah Matriks Kiri Menjadi Identitas
    doc.add_heading('Langkah 2: Operasi Baris', level=1)
    
    # Sublangkah 2.1: Bagi Baris Pertama dengan 2
    doc.add_heading('2.1. Bagi Baris Pertama dengan 2', level=2)
    table_2_1 = doc.add_table(rows=3, cols=6)
    data_2_1 = [
        [1, 0, 0, 1/2, 0, 0],
        [1, 5, 0, 0, 1, 0],
        [7, 9, 4, 0, 0, 1]
    ]
    for i in range(3):
        row = data_2_1[i]
        for j in range(6):
            table_2_1.cell(i, j).text = str(row[j])
    
    # Sublangkah 2.2: Kurangi Baris Kedua dengan Baris Pertama
    doc.add_heading('2.2. Kurangi Baris Kedua dengan Baris Pertama', level=2)
    table_2_2 = doc.add_table(rows=3, cols=6)
    data_2_2 = [
        [1, 0, 0, 1/2, 0, 0],
        [0, 5, 0, -1/2, 1, 0],
        [7, 9, 4, 0, 0, 1]
    ]
    for i in range(3):
        row = data_2_2[i]
        for j in range(6):
            table_2_2.cell(i, j).text = str(row[j])
    
    # Sublangkah 2.3: Kurangi 7xBaris Pertama dari Baris Ketiga
    doc.add_heading('2.3. Kurangi 7xBaris Pertama dari Baris Ketiga', level=2)
    table_2_3 = doc.add_table(rows=3, cols=6)
    data_2_3 = [
        [1, 0, 0, 0.5, 0, 0],
        [0, 5, 0, -0.5, 1, 0],
        [0, 9, 4, -3.5, 0, 1]
    ]
    for i in range(3):
        row = data_2_3[i]
        for j in range(6):
            table_2_3.cell(i, j).text = str(row[j])
    
    # Sublangkah 2.4: Bagi Baris Kedua dengan 5
    doc.add_heading('2.4. Bagi Baris Kedua dengan 5', level=2)
    table_2_4 = doc.add_table(rows=3, cols=6)
    data_2_4 = [
        [1, 0, 0, 0.5, 0, 0],
        [0, 1, 0, -1/10, 1/5, 0],
        [0, 9, 4, -3.5, 0, 1]
    ]
    for i in range(3):
        row = data_2_4[i]
        for j in range(6):
            table_2_4.cell(i, j).text = str(row[j])
    
    # Sublangkah 2.5: Kurangi 9xBaris Kedua dari Baris Ketiga
    doc.add_heading('2.5. Kurangi 9xBaris Kedua dari Baris Ketiga', level=2)
    table_2_5 = doc.add_table(rows=3, cols=6)
    data_2_5 = [
        [1, 0, 0, 0.5, 0, 0],
        [0, 1, 0, -0.1, 0.2, 0],
        [0, 0, 4, -16/5, -9/5, 1]
    ]
    for i in range(3):
        row = data_2_5[i]
        for j in range(6):
            table_2_5.cell(i, j).text = str(row[j])
    
    # Sublangkah 2.6: Bagi Baris Ketiga dengan 4
    doc.add_heading('2.6. Bagi Baris Ketiga dengan 4', level=2)
    table_2_6 = doc.add_table(rows=3, cols=6)
    data_2_6 = [
        [1, 0, 0, 0.5, 0, 0],
        [0, 1, 0, -0.1, 0.2, 0],
        [0, 0, 1, -4/5, -9/20, 1/4]
    ]
    for i in range(3):
        row = data_2_6[i]
        for j in range(6):
            table_2_6.cell(i, j).text = str(row[j])
    
    # Hasil Akhir: Invers Matriks
    doc.add_heading('Hasil Akhir: Invers Matriks \( A^{-1} \)', level=1)
    inverse_table = doc.add_table(rows=3, cols=3)
    inverse_data = [
        [0.5, 0, 0],
        [-1/10, 1/5, 0],
        [-4/5, -9/20, 0.25]
    ]
    for i in range(3):
        row = inverse_data[i]
        for j in range(3):
            inverse_table.cell(i, j).text = f"{row[j]:.4f}"
    
    # Simpan dokumen
    doc.save('invers_matriks.docx')
    print("Dokumen berhasil dibuat dan disimpan sebagai 'invers_matriks.docx'.")
    
if __name__ == "__main__":
    create_document()